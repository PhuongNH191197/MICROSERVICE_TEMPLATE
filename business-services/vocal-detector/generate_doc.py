"""Generate VOCAL_DETECTOR_API.docx — run via Docker, output to /workspace."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUT = "/workspace/VOCAL_DETECTOR_API.docx"
doc = Document()

# ── page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── helpers ───────────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def h3(text):
    doc.add_heading(text, level=3)

def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11)

def bullet(text):
    doc.add_paragraph(text, style="List Bullet")

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "2E74B5")
        tcPr.append(shd)
    for r_idx, row_data in enumerate(rows):
        row = t.rows[r_idx + 1]
        fill = "FFFFFF" if r_idx % 2 == 0 else "EBF3FB"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  fill)
            tcPr.append(shd)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("VOCAL DETECTOR SERVICE")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("API Documentation & Integration Guide")
r2.font.size = Pt(14)
r2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
r2.italic = True

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run(f"Version 1.0  |  {datetime.date.today().strftime('%B %d, %Y')}")
r3.font.size = Pt(11)
r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
h1("1. Service Overview")
body(
    "The Vocal Detector Service is a RESTful microservice that analyzes audio files "
    "to determine whether they contain human vocals or are purely instrumental. "
    "It is designed for integration into music processing pipelines, content moderation "
    "systems, or any application that needs to classify audio tracks automatically."
)
doc.add_paragraph()
h2("Key Features")
bullet("Supports MP3, WAV, FLAC, M4A input formats")
bullet("Maximum file size: 50 MB per request")
bullet("Analyzes only the first 60 seconds (optimized for CPU-only environments)")
bullet("Returns probability scores for both vocal and instrumental classes")
bullet("Single HTTP endpoint — simple to integrate")
bullet("Containerized with Docker — zero dependency installation on the host")
doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# 2. TECH STACK
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Technology Stack")
add_table(
    ["Component", "Technology", "Version", "Role"],
    [
        ["Runtime",           "Python",               "3.11",            "Application language"],
        ["Web Framework",     "FastAPI",               ">= 0.110",        "HTTP API, async request handling"],
        ["ASGI Server",       "Uvicorn",               ">= 0.29",         "Production-grade ASGI server"],
        ["ML / DSP Library",  "Essentia-TensorFlow",   "Latest",          "Audio feature extraction & inference"],
        ["Deep Learning",     "TensorFlow",            "2.x (CPU)",       "Bundled with Essentia-TF"],
        ["Audio Conversion",  "FFmpeg",                "System",          "Decode MP3 / FLAC / M4A to WAV"],
        ["Containerization",  "Docker",                ">= 24",           "Isolated runtime environment"],
        ["Base Image",        "python:3.11-slim",      "Debian Bookworm", "Minimal Linux base"],
    ],
    [4, 4.5, 3, 6]
)


# ══════════════════════════════════════════════════════════════════════════════
# 3. AI MODELS
# ══════════════════════════════════════════════════════════════════════════════
h1("3. AI Model Information")
body(
    "The service uses two pre-trained neural network models published by the "
    "Music Technology Group (MTG) at Universitat Pompeu Fabra, Barcelona. "
    "Both models are downloaded from the official MTG model repository (essentia.upf.edu)."
)
doc.add_paragraph()
add_table(
    ["Model", "File", "Size", "Purpose"],
    [
        ["Discogs-EffNet Embedder",
         "discogs-effnet-bs64-1.pb",
         "17.5 MB",
         "Extracts 1280-dim audio embeddings. EfficientNet-B0 trained on Discogs dataset."],
        ["Voice/Instrumental Classifier",
         "voice_instrumental-discogs-effnet-1.pb",
         "502 KB",
         "Binary classifier on EffNet embeddings. Outputs [instrumental, voice] probabilities."],
        ["Classifier Metadata",
         "voice_instrumental-discogs-effnet-1.json",
         "< 1 KB",
         "Maps output indices to class labels (instrumental / voice)."],
    ],
    [4.5, 5.5, 2, 5.5]
)
h2("Inference Pipeline")
for step in [
    "1. Upload received and validated (format + size check)",
    "2. FFmpeg converts input to 16 kHz mono WAV, trimmed to 60 seconds",
    "3. Essentia MonoLoader reads WAV into float32 audio array",
    "4. TensorflowPredictEffnetDiscogs extracts per-frame embeddings",
    "5. TensorflowPredict2D classifies embeddings → per-frame [instrumental, voice] scores",
    "6. Scores averaged across all frames → final probabilities",
    "7. JSON response returned to caller",
]:
    bullet(step)
doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# 4. SYSTEM REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════════
h1("4. System Requirements")
add_table(
    ["Requirement", "Minimum", "Recommended"],
    [
        ["CPU",    "2 cores",   "4 cores"],
        ["RAM",    "1 GB",      "2 GB"],
        ["Disk",   "3 GB",      "5 GB (Docker image ~2.5 GB)"],
        ["Docker", "24.x",      "Latest stable"],
        ["OS",     "Linux / Windows 10+ / macOS", "Linux (production)"],
        ["GPU",    "Not required", "Not supported (CPU-only TF)"],
    ],
    [4, 5, 8.5]
)


# ══════════════════════════════════════════════════════════════════════════════
# 5. INSTALLATION
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Installation & Setup")

h2("Step 1 — Obtain the service files")
body("The service directory must have the following structure:")
code(
    "vocal-detector/\n"
    "├── app/\n"
    "│   ├── main.py\n"
    "│   ├── detector.py\n"
    "│   └── schemas.py\n"
    "├── models/          <-- model files go here (Step 2)\n"
    "├── scripts/\n"
    "│   ├── download_models.ps1\n"
    "│   └── download_models.sh\n"
    "├── Dockerfile\n"
    "├── docker-compose.yml\n"
    "└── requirements.txt"
)

h2("Step 2 — Download AI model files")
body("Windows (PowerShell):")
code("cd vocal-detector\n.\\scripts\\download_models.ps1")
body("Linux / macOS:")
code("cd vocal-detector\nbash scripts/download_models.sh")
body("Downloads 3 files (~18 MB total) into models/. Internet required for this step only.")

h2("Step 3 — Build and start")
body("Windows — run from PowerShell (NOT Git Bash — it mangles container paths):")
code("cd vocal-detector\ndocker-compose up -d --build")
body("First build: ~5-10 minutes. Subsequent starts: <30 seconds.")

h2("Step 4 — Verify")
code("curl http://localhost:8765/health")
body('Expected: {"status": "ok"}')
doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# 6. API REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
h1("6. API Reference")

h2("Base URL")
code("http://<host>:8765")

h2("GET  /health")
body("Health check. Returns 200 when service is ready to accept requests.")
add_table(
    ["Property", "Value"],
    [
        ["Method",         "GET"],
        ["URL",            "/health"],
        ["Auth",           "None"],
        ["Response 200",   '{"status": "ok"}'],
    ],
    [4, 13.5]
)

h2("POST  /detect")
body("Analyze an audio file. Returns vocal detection result.")
add_table(
    ["Property", "Value"],
    [
        ["Method",           "POST"],
        ["URL",              "/detect"],
        ["Content-Type",     "multipart/form-data"],
        ["Auth",             "None"],
        ["Max file size",    "50 MB"],
        ["Accepted formats", "mp3, wav, flac, m4a"],
    ],
    [4, 13.5]
)

h3("Request Fields")
add_table(
    ["Field", "Type", "Required", "Description"],
    [["file", "binary", "Yes", "Audio file to analyze"]],
    [3, 3.5, 3, 8]
)

h3("Response Body — 200 OK")
code(
    "{\n"
    '  "is_instrumental": true,\n'
    '  "confidence": 0.9601,\n'
    '  "scores": {\n'
    '    "instrumental": 0.9601,\n'
    '    "voice": 0.0399\n'
    "  }\n"
    "}"
)

h3("Response Fields")
add_table(
    ["Field", "Type", "Description"],
    [
        ["is_instrumental",      "boolean",    "true = no vocals detected. false = vocals present."],
        ["confidence",           "float 0–1",  "Certainty of the prediction. = max(scores.instrumental, scores.voice)."],
        ["scores.instrumental",  "float 0–1",  "Probability the audio is instrumental."],
        ["scores.voice",         "float 0–1",  "Probability the audio contains vocals."],
    ],
    [4, 3, 10.5]
)
body("scores.instrumental + scores.voice = 1.0  (Softmax output)")

h3("Error Codes")
add_table(
    ["HTTP Status", "Trigger", "Response"],
    [
        ["400", "Unsupported file format",    '{"detail": "Unsupported format..."}'],
        ["413", "File > 50 MB",               '{"detail": "File exceeds 50 MB limit."}'],
        ["422", "Corrupt / unreadable audio", '{"detail": "Audio decoding failed."}'],
        ["500", "Unexpected server error",    '{"detail": "Internal server error"}'],
    ],
    [2.5, 5, 10]
)


# ══════════════════════════════════════════════════════════════════════════════
# 7. USAGE EXAMPLES
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Usage Examples")

h2("cURL — Linux / macOS / Windows CMD")
code(
    "curl -X POST http://localhost:8765/detect \\\n"
    '  -F "file=@/path/to/song.mp3"'
)

h2("cURL — Windows PowerShell")
body("Use curl.exe (not curl) to bypass the PowerShell alias:")
code(
    'curl.exe -X POST http://localhost:8765/detect `\n'
    '  -F "file=@C:\\path\\to\\song.mp3"'
)

h2("Python — requests")
code(
    "import requests\n\n"
    'with open("song.mp3", "rb") as f:\n'
    '    resp = requests.post(\n'
    '        "http://localhost:8765/detect",\n'
    '        files={"file": ("song.mp3", f, "audio/mpeg")},\n'
    "    )\n\n"
    "data = resp.json()\n"
    'print(data["is_instrumental"], data["confidence"])'
)

h2("JavaScript — fetch")
code(
    "const form = new FormData();\n"
    'form.append("file", audioFile);\n\n'
    'const resp = await fetch("http://localhost:8765/detect", {\n'
    '  method: "POST", body: form\n'
    "});\n"
    "const data = await resp.json();\n"
    "console.log(data.is_instrumental, data.confidence);"
)
doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# 8. CONFIGURATION
# ══════════════════════════════════════════════════════════════════════════════
h1("8. Configuration")
add_table(
    ["Environment Variable", "Default", "Description"],
    [
        ["MODELS_DIR", "/models",
         "Path inside the container where model files are mounted. Change only if volume mount path is also changed."],
    ],
    [5, 3, 9.5]
)
body("Resource limits (docker-compose.yml):")
code(
    "deploy:\n"
    "  resources:\n"
    "    limits:\n"
    "      cpus: '2'\n"
    "      memory: 1G"
)
doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# 9. LIMITATIONS
# ══════════════════════════════════════════════════════════════════════════════
h1("9. Known Limitations")
add_table(
    ["Limitation", "Details"],
    [
        ["Audio length",    "Only the first 60 seconds analyzed. Files < 5s may be inaccurate."],
        ["Mixed content",   "Tracks mixing instrumental + short vocal phrases may yield low confidence (~0.5)."],
        ["File formats",    "Only mp3, wav, flac, m4a. Other formats return HTTP 400."],
        ["File size",       "Max 50 MB. Trim large files before sending."],
        ["Concurrency",     "Single worker. High-throughput deployments should run multiple container replicas behind a load balancer."],
        ["GPU",             "Not supported. Inference runs on CPU only."],
    ],
    [4, 13.5]
)


# ══════════════════════════════════════════════════════════════════════════════
# 10. TROUBLESHOOTING
# ══════════════════════════════════════════════════════════════════════════════
h1("10. Troubleshooting")
add_table(
    ["Symptom", "Cause", "Fix"],
    [
        ["Model file not found at startup",
         "models/ directory empty or volume not mounted",
         "Run download script, then start with docker-compose from PowerShell"],
        ["curl fails in PowerShell",
         "PS aliases curl to Invoke-WebRequest",
         "Use curl.exe instead"],
        ["CUDA / libcudart warnings in logs",
         "No GPU; TF looks for CUDA on init",
         "Harmless — ignore. Service runs on CPU."],
        ["HTTP 413",
         "File > 50 MB",
         "Compress or trim the audio file"],
        ["HTTP 422",
         "Corrupt audio file",
         "Verify file plays in a media player"],
        ["Service starts but returns low confidence (0.4-0.6)",
         "Track has mixed vocals/instrumental sections",
         "Expected behavior — use a threshold appropriate for your use case"],
    ],
    [4.5, 4.5, 8.5]
)


# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
    "Service version: 1.0.0\n"
    "AI Model source: Music Technology Group (MTG) — Universitat Pompeu Fabra\n"
    "essentia.upf.edu"
)
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
r.italic = True

doc.save(OUT)
print(f"Saved: {OUT}")
